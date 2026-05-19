"""Extract text from PDF and Word documents for contract review."""

from __future__ import annotations

from pathlib import Path


def extract_text(file_path: str, filename: str = "") -> tuple[str, str]:
    """Extract text from a file. Returns (text, error_message).

    Supported formats: .pdf, .docx, .txt, .md
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    if not suffix:
        # Try to detect from provided filename
        suffix = Path(filename).suffix.lower() if filename else ""

    try:
        if suffix in (".txt", ".md", ""):
            return _read_txt(path), ""
        elif suffix == ".pdf":
            return _read_pdf(path), ""
        elif suffix in (".docx", ".doc"):
            return _read_docx(path), ""
        else:
            return "", f"不支持的文件格式: {suffix}"
    except Exception as e:
        return "", f"文件解析失败: {e}"


def _read_txt(path: Path) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(path, encoding=enc) as f:
                return f.read().strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read().strip()


def _read_pdf(path: Path) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    result = "\n\n".join(pages)
    if result.strip():
        return result
    # Text layer empty — try OCR for scanned/image PDFs
    return _read_pdf_ocr(path)


def _read_pdf_ocr(path: Path) -> str:
    """OCR fallback for scanned/image-based PDFs using pypdfium2 + Tesseract."""
    import pypdfium2 as pdfium

    try:
        import pytesseract
    except ImportError:
        raise ValueError(
            "PDF 为扫描件且未安装 OCR 库（pytesseract）。"
            " 请运行: pip install pytesseract"
        )

    # Point to Tesseract installation
    pytesseract.pytesseract.tesseract_cmd = r"E:\Tesseract-OCR\tesseract.exe"

    pdf = pdfium.PdfDocument(str(path))
    n_pages = len(pdf)
    pages: list[str] = []
    try:
        for i in range(n_pages):
            page = pdf[i]
            bitmap = page.render(scale=2)
            pil_image = bitmap.to_pil()
            text = pytesseract.image_to_string(pil_image, lang="chi_sim+eng")
            if text:
                pages.append(text.strip())
    finally:
        pdf.close()

    result = "\n\n".join(pages)
    if not result.strip():
        raise ValueError(
            f"PDF 为扫描件（共 {n_pages} 页），OCR 未识别到文字。"
            " 请确认已安装 Tesseract OCR 引擎及中文语言包。"
        )
    return result


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    result = "\n".join(paragraphs)
    if not result.strip():
        raise ValueError("Word 文档内容为空")
    return result
